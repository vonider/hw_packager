from docx import Document
from docx.shared import Inches, Mm
from datetime import datetime
import os

f = open('config.txt', 'r', encoding='utf-8')
student = f.readline()[13:].strip('\n').replace('"', '')
group = f.readline()[6:].replace('"', '')

# Find photos
photos_dir = './photos/'
files = os.listdir(photos_dir)

# Set size
document = Document()
section = document.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)

# Metadata changing
core_properties = document.core_properties
core_properties.author = ''
core_properties.comments = ''

# Data for filename and document
subject = input("Работа по дисциплине: ")
task = input("Задание: ").capitalize()
date = str(datetime.now().strftime('%d-%m-%y'))

# Creating file
document.add_heading(f'Работа по дисциплине: {subject}', 0)
doc = document.add_paragraph(f'Выполнил(а): {student}\n')
doc.add_run(f'Группа: {group}')

for file in files:
    document.add_picture(f'{photos_dir}{file}', width=Inches(6), height=Inches(8))

# Saving
document.save(f'{student} {group} {subject} {task} {date}.docx')
try:
    os.remove('program.py.bak')
except:
    pass